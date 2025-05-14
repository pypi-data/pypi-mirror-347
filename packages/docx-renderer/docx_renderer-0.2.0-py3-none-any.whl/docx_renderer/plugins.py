from __future__ import annotations
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt
from .exceptions import RenderError
from .utils import convert_to_length

def image(
    context: dict,
    width: int | str | None = None,
    height: int | str | None = None,
    remove_placeholder: bool = True,
):
    """Insert an image into the document at the location of the placeholder.

    Args:
        context (dict): Dictionary containing the following keys:
            - result: The result of evaluating the python statement.
            - paragraph: The docx paragraph object where the placeholder is present.
            - document: The output docx document object.
        width (int | str | None, optional): Width of the image. Can be an int or a string with a unit.
            Defaults to None.
        height (int | str | None, optional): Height of the image. Can be an int or a string with a unit.
            Defaults to None.
        remove_placeholder (bool, optional): Remove the placeholder after the image is inserted.
            Defaults to True.

    Raises:
        RenderError: If the image file does not exist.
    """
    result = str(context["result"])
    container = context["paragraph"]
    document = context["document"]

    if not Path(result).exists():
        raise RenderError(f"Image '{result}' not found.")

    if remove_placeholder:
        container.text = ""
    
    width = convert_to_length(width) if width else None
    height = convert_to_length(height) if height else None

    container.add_run().add_picture(result, width=width, height=height)

    # Optionally remove the placeholder text
def table(
    context: dict,
    style: str | None = None,
    remove_placeholder: bool = True,
):
    """Insert a table into the document at the location of the placeholder.

    Args:
        context (dict): Dictionary containing the following keys:
            - result: The result of evaluating the python statement.
            - paragraph: The docx paragraph object where the placeholder is present.
            - document: The output docx document object.
        first_row (bool, optional): Show the first row as header. Defaults to True.
        remove_placeholder (bool, optional): Remove the placeholder after the table is inserted.
            Defaults to True.
    """
    result = context["result"]
    container = context["container"]
    document = context["document"]

    table_data = list(result)
    table = document.add_table(rows=len(table_data), cols=len(table_data[0]), style=style)

    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            row.cells[col_idx].text = str(cell_data)
    tbl, p = table._tbl, container._p
    p.addnext(tbl)

    # Optionally remove the placeholder text
    if remove_placeholder:
        container.text = ""
