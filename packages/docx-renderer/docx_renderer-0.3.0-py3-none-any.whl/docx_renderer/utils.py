from docx.shared import Inches, Pt, Cm, Mm, Length

def container_text_replace(obj, find_string, replace_string):
    """Function to replace text in a paragraph

    This function replaces text in a paragraph while respecting the formatting.

    Args:
        para: Paragraph/cell to replace text in.
        find_string (str): String to find in the paragraph.
        replace_string (str): String to replace the find_string with.

    Returns:
        None
    """
    find_string = str(find_string)
    replace_string = str(replace_string)
    starting_pos = obj.text.find(find_string)
    if starting_pos == -1:
        return  # text not in paragraph
    txt_prev = ""
    for run in obj.runs:
        if len(txt_prev) <= starting_pos < len(txt_prev) + len(run.text):
            if run.text.find(find_string) != -1:  # text in run, replace
                run.text = run.text.replace(find_string, replace_string)
                return
            else:  # text no in "run"
                txt_prev = txt_prev + run.text
                run.text = run.text[: starting_pos - len(txt_prev)] + replace_string
        elif starting_pos < len(txt_prev) and starting_pos + len(find_string) >= len(
            txt_prev
        ) + len(run.text):
            txt_prev = txt_prev + run.text
            run.text = ""
        elif (
            len(txt_prev)
            < starting_pos + len(find_string)
            < len(txt_prev) + len(run.text)
        ):
            txt_prev = txt_prev + run.text
            run.text = run.text[starting_pos + len(find_string) - len(txt_prev) :]
        else:
            txt_prev += run.text


def fix_quotes(input_string: str) -> str:
    """Replace unicode quotes (inserted by powerpoint) with ascii quotes.

    Args:
        input_string (str): String to fix quotes in.
    
    Returns:
        str: String with fixed quotes.
    """
    return (
        input_string.replace("’", "'")
        .replace("‘", "'")
        .replace("“", '"')
        .replace("”", '"')
    )

def copy_table(source_shape, target_slide):
    source_table = source_shape.table
    # Get the dimensions of the source table
    rows = len(source_table.rows)
    cols = len(source_table.columns)

    # Add a new table to the target slide with the same dimensions
    left = source_shape.left
    top = source_shape.top
    width = source_shape.width
    height = source_shape.height
    target_table = target_slide.shapes.add_table(
        rows, cols, left, top, width, height
    ).table

    # Copy the content and formatting of each cell
    for row_idx in range(rows):
        for col_idx in range(cols):
            source_cell = source_table.cell(row_idx, col_idx)
            target_cell = target_table.cell(row_idx, col_idx)

            copy_text_frame(source_cell.text_frame, target_cell.text_frame)


def copy_text_frame(source_text_frame, target_text_frame):
    try:
        target_text_frame.clear()  # Clear any existing text
    except AttributeError:
        pass  # Text frame is not clearable
    if not hasattr(source_text_frame, "paragraphs"):
        return
    for i, paragraph in enumerate(source_text_frame.paragraphs):
        if i == 0:
            # First paragraph is already there
            new_paragraph = target_text_frame.paragraphs[0]
        else:
            new_paragraph = target_text_frame.add_paragraph()
        new_paragraph.level = paragraph.level

        # Copy paragraph formatting
        new_paragraph.font.bold = paragraph.font.bold
        new_paragraph.font.italic = paragraph.font.italic
        new_paragraph.font.underline = paragraph.font.underline
        new_paragraph.font.size = paragraph.font.size
        if paragraph.font.color and hasattr(paragraph.font.color, "rgb"):
            new_paragraph.font.color.rgb = paragraph.font.color.rgb
        new_paragraph.alignment = paragraph.alignment

        for run in paragraph.runs:
            new_run = new_paragraph.add_run()
            new_run.text = run.text

            # Copy run formatting
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            new_run.font.underline = run.font.underline
            new_run.font.size = run.font.size
            if run.font.color and hasattr(run.font.color, "rgb") and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            elif run.font.color and hasattr(run.font.color, "theme_color"):
                new_run.font.color.theme_color = run.font.color.theme_color
            if run.font.color and hasattr(run.font.color, "brightness") and run.font.color.brightness:
                new_run.font.color.brightness = run.font.color.brightness
            if (
                hasattr(run, "hyperlink")
                and run.hyperlink
                and hasattr(run.hyperlink, "address")
            ):
                new_run.hyperlink.address = run.hyperlink.address

def convert_to_length(length: int | str) -> Length:
    """Convert a length to a Length object.

    Args:
        length (int | str): Length to convert. Can be an int or a string with a unit.

    Returns:
        Length: Length object.
    """
    if isinstance(length, str):
        if length.endswith("cm"):
            return Cm(int(length[:-2]))
        elif length.endswith("in"):
            return Inches(int(length[:-2]))
        elif length.endswith("pt"):
            return Pt(int(length[:-2]))
        elif length.endswith("mm"):
            return Mm(int(length[:-2]))
        else:
            raise ValueError(f"Invalid length format: {length}")
    elif isinstance(length, int):
        return Length(length)
    else:
        raise TypeError(f"Invalid type for length: {type(length)}")