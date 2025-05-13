import os
import re
import warnings
from typing import Dict, List, Optional

import evaluate
import numpy as np
import pandas as pd
import torch
import tqdm
from datasets import Dataset
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from sklearn.model_selection import train_test_split
from transformers import (
    AutoModelForSequenceClassification,
    AutoTokenizer,
    DataCollatorWithPadding,
    Trainer,
    TrainingArguments,
    pipeline,
)

warnings.filterwarnings("ignore")


class TranscriptTrainer:
    def __init__(
        self,
        input_file: str,
        destination_path: Optional[str] = None,
        text_column: str = "full_quote",
        target_column: str = "target",
        test_size: int = 0.2,
        max_length: int = 512,
        num_train_epochs: int = 20,
        learning_rate_distilbert: float = 1e-5,
        learning_rate_electra: float = 2e-5,
        labels: Optional[List[str]] = None,
        upper_lower_mapping: Optional[Dict[str, List[str]]] = None,
    ):
        """
        Initialize the TranscriptTrainer

        Args:
            input_file (str): Path to the input file (.csv or .xlsx format)
            destination_path (str, optional): Path to save model outputs and results
            max_length (int, optional): Maximum sequence length for tokenization
            num_train_epochs (int, optional): Number of training epochs
            learning_rate_distilbert (float, optional): Learning rate for DistilBERT
            learning_rate_electra (float, optional): Learning rate for Electra
            labels (List[str], optional): List of classification labels
            upper_lower_mapping (Dict[str, List[str]], optional): Dictionary mapping top-level categories to their subcategories
        """
        # Default Labels
        self.LABELS = labels or [
            "Value equation",
            "Credentialing / Quality Assurance Infrastructure",
            "Financial Impact",
            "Health System Characteristics",
            "Clinical utility & efficiency-Provider perspective",
            "Workflow related problems",
            "Provider Characteristics",
            "Training",
            "Patient/Physician interaction in LUS",
            "Imaging modalities in general",
        ]

        # Default Top Levels (with more flexible structure)
        self.upper_lower_mapping = upper_lower_mapping or {
            "multi_level_org_char": [
                "Provider Characteristics",
                "Health System Characteristics",
            ],
            "multi_level_org_perspect": [
                "Imaging modalities in general",
                "Value equation",
                "Clinical utility & efficiency-Provider perspective",
                "Patient/Physician interaction in LUS",
                "Workflow related problems",
            ],
            "impl_sust_infra": [
                "Training",
                "Credentialing / Quality Assurance Infrastructure",
                "Financial Impact",
            ],
        }

        # Flatten categories for easier processing
        self.CATEGORIES = {}
        for top_level, subcategories in self.upper_lower_mapping.items():
            for category in subcategories:
                self.CATEGORIES[category] = top_level

        # Configurations
        self.input_file = input_file
        self.destination_path = destination_path or os.path.dirname(input_file)
        self.max_length = max_length
        self.num_train_epochs = num_train_epochs
        self.learning_rate_distilbert = learning_rate_distilbert
        self.learning_rate_electra = learning_rate_electra
        self.test_size = test_size
        self.text_column = text_column

        self.bert_checkpoint = None
        self.electra_checkpoint = None

        # Devices
        self.device = "cuda" if torch.cuda.is_available() else "cpu"

        # Load and prepare data
        self._load_and_prepare_data()
        print("DATASET LABELS AND IDS PREPARED SUCCESSFULLY")

    def _load_and_prepare_data(self):
        """Load input data and prepare for training"""
        # Determine file type and load accordingly
        file_extension = os.path.splitext(self.input_file)[1].lower()

        if file_extension == ".csv":
            full_data = pd.read_csv(self.input_file)
        elif file_extension == ".xlsx":
            full_data = pd.read_excel(self.input_file)
        else:
            raise ValueError(
                f"Unsupported file format: {file_extension}. Please use .csv or .xlsx files."
            )

        # Reset index to ensure unique labels
        full_data = full_data.reset_index(drop=True)

        # Add target column
        full_data["target"] = "Unknown"

        # Assign target based on label columns
        for col in self.LABELS:
            full_data.loc[full_data[col] == 1, "target"] = col

        # Perform train-test split
        msk_train, msk_test = train_test_split(
            full_data.index, test_size=0.2, stratify=full_data.target, random_state=0
        )

        # Use vectorized operations to assign split labels
        full_data["split"] = "neither"
        full_data.loc[msk_train, "split"] = "train"
        full_data.loc[msk_test, "split"] = "test"

        # Remove Unknown target rows
        full_data = full_data[full_data["target"] != "Unknown"]

        # Add numeric labels
        full_data["labels"] = full_data.target.astype("category").cat.codes

        # Save to CSV instead of Excel
        full_data.to_csv(
            os.path.join(self.destination_path, "SplitHuman.csv"), index=False
        )

        # Store the processed dataframe
        self.full_data = full_data

        # Generate label mappings
        labels = full_data.target.unique().tolist()
        self.label2id = {i: labels.index(i) for i in labels}
        self.id2label = {v: k for k, v in self.label2id.items()}

        # Prepare train and eval datasets
        self.train_data = full_data[full_data["split"] == "train"]
        self.eval_data = full_data[full_data["split"] == "test"]

        # Create Hugging Face datasets

        # USING THE FULL DATA TO TRAIN, THIS WOULD BE CHANGED SOON...
        self.train_dataset = Dataset.from_dict(
            {
                "text": full_data[self.text_column].values.tolist(),
                "label": full_data.labels.tolist(),
            }
        )
        self.eval_dataset = Dataset.from_dict(
            {
                "text": self.eval_data[self.text_column].values.tolist(),
                "label": self.eval_data.labels.tolist(),
            }
        )

        # Print dataset shape for verification
        print(f"Full dataset shape: {full_data.shape}")

    def generate_upper_level_columns(self):
        """
        Generate upper level columns based on upper_lower_mapping.
        Places new columns before conventional label columns.
        """

        # Create new columns for each upper level category
        for upper_category, lower_categories in self.upper_lower_mapping.items():
            # Initialize the new column with False
            self.full_data[upper_category] = False

            # Set True if any of the lower categories are True
            for lower_cat in lower_categories:
                if lower_cat in self.full_data.columns:
                    self.full_data[upper_category] |= self.full_data[lower_cat]

        # Reorder columns to place upper level columns before conventional labels
        # Get all columns except the new upper level columns
        existing_cols = [
            col
            for col in self.full_data.columns
            if col not in self.upper_lower_mapping.keys()
        ]

        # Find the position of first label column
        label_start_idx = next(
            i for i, col in enumerate(existing_cols) if col in self.LABELS
        )

        # Reorder columns
        new_column_order = (
            existing_cols[:label_start_idx]
            + list(self.upper_lower_mapping.keys())
            + existing_cols[label_start_idx:]
        )

        # Apply new column order
        self.full_data = self.full_data[new_column_order]

    def _preprocess_data(self, tokenizer):
        """Preprocess data for model training"""

        def preprocess_function(examples):
            return tokenizer(
                examples["text"], truncation=True, max_length=self.max_length
            )

        train_dataset = self.train_dataset.map(preprocess_function)
        eval_dataset = self.eval_dataset.map(preprocess_function)

        return train_dataset, eval_dataset

    def _compute_metrics(self, eval_pred):
        """Compute accuracy metrics"""
        accuracy = evaluate.load("accuracy")
        predictions, labels = eval_pred
        predictions = np.argmax(predictions, axis=1)
        return accuracy.compute(predictions=predictions, references=labels)

        print("COMPUTE METRIC LOADED")

    def train_distilbert(self):
        """Train DistilBERT model"""
        output_dir = os.path.join(self.destination_path, "bert_weights")
        os.makedirs(output_dir, exist_ok=True)

        model = AutoModelForSequenceClassification.from_pretrained(
            "distilbert-base-uncased", id2label=self.id2label, label2id=self.label2id
        ).to(self.device)

        tokenizer = AutoTokenizer.from_pretrained(
            "distilbert-base-uncased", max_length=self.max_length
        )

        train_dataset, eval_dataset = self._preprocess_data(tokenizer)
        data_collator = DataCollatorWithPadding(tokenizer=tokenizer)

        training_args = TrainingArguments(
            output_dir=output_dir,
            learning_rate=self.learning_rate_distilbert,
            per_device_train_batch_size=2,
            per_device_eval_batch_size=2,
            num_train_epochs=self.num_train_epochs,
            weight_decay=0.01,
            evaluation_strategy="epoch",
            save_strategy="epoch",
            load_best_model_at_end=True,
            report_to=["none"],
        )

        trainer = Trainer(
            model=model,
            args=training_args,
            train_dataset=train_dataset,
            eval_dataset=eval_dataset,
            tokenizer=tokenizer,
            data_collator=data_collator,
            compute_metrics=self._compute_metrics,
        )

        trainer.train()

        return trainer

    def train_electra(self):
        """Train Electra model"""
        output_dir = os.path.join(self.destination_path, "electra_weights")
        os.makedirs(output_dir, exist_ok=True)

        model = AutoModelForSequenceClassification.from_pretrained(
            "mrm8488/electra-small-finetuned-squadv2",
            id2label=self.id2label,
            label2id=self.label2id,
        ).to(self.device)

        tokenizer = AutoTokenizer.from_pretrained(
            "mrm8488/electra-small-finetuned-squadv2", max_length=self.max_length
        )

        train_dataset, eval_dataset = self._preprocess_data(tokenizer)
        data_collator = DataCollatorWithPadding(tokenizer=tokenizer)

        training_args = TrainingArguments(
            output_dir=output_dir,
            learning_rate=self.learning_rate_electra,
            per_device_train_batch_size=2,
            per_device_eval_batch_size=2,
            num_train_epochs=self.num_train_epochs,
            weight_decay=0.01,
            evaluation_strategy="epoch",
            save_strategy="epoch",
            load_best_model_at_end=True,
            report_to=["none"],
        )

        trainer = Trainer(
            model=model,
            args=training_args,
            train_dataset=train_dataset,
            eval_dataset=eval_dataset,
            tokenizer=tokenizer,
            data_collator=data_collator,
            compute_metrics=self._compute_metrics,
        )

        trainer.train()

        return trainer

    def _classify_text(self, text, pipeline_model):
        """
        Classify a single text using a pipeline model

        Args:
            text (str): Text to classify
            pipeline_model (pipeline): Trained pipeline model

        Returns:
            dict: Classification result with additional top-level information
        """
        output = pipeline_model(text, max_length=self.max_length, truncation=True)

        # Add top-level category to the output
        label = output[0]["label"]
        output[0]["top level"] = self.CATEGORIES.get(label, "")

        return output

    def classify_sheet_with_model(self, model_path, input_dataframe):
        """
        Classify text using a trained model

        Args:
            model_path (str): Path to the trained model
            input_dataframe (pd.DataFrame): DataFrame to classify

        Returns:
            pd.DataFrame: Classified DataFrame
        """
        model_pipeline = pipeline("text-classification", model_path)
        result_sheet = input_dataframe.copy()

        for index, row in result_sheet.iterrows():
            if row[self.LABELS].values.sum() > 0:
                prediction = self._classify_text(row[self.text_column], model_pipeline)
                feature = prediction[0]["label"]

                result_sheet.at[index, feature] = 1

                # Add top-level category columns
                for top_level, categories in self.upper_lower_mapping.items():
                    column_name = top_level.lower().replace(" ", "_")
                    if feature in categories:
                        result_sheet.at[index, column_name] = 1

        return result_sheet

    def _get_latest_checkpoint(self, model_dir):
        """
        Get the latest checkpoint from a model directory

        Args:
            model_dir (str): Directory containing model checkpoints

        Returns:
            str: Path to the latest checkpoint
        """
        checkpoints = [
            os.path.join(model_dir, d)
            for d in os.listdir(model_dir)
            if d.startswith("checkpoint-")
        ]
        latest_checkpoint = max(checkpoints, key=lambda x: int(x.split("-")[-1]))
        return latest_checkpoint

    def train_and_classify(self):
        """Train both models and classify"""
        bert_trainer = self.train_distilbert()
        electra_trainer = self.train_electra()

        # Use the best model checkpoint rather than the latest checkpoint.
        self.bert_checkpoint = bert_trainer.state.best_model_checkpoint or os.path.join(
            self.destination_path, "bert_weights"
        )
        self.electra_checkpoint = (
            electra_trainer.state.best_model_checkpoint
            or os.path.join(self.destination_path, "electra_weights")
        )

        bert_sheet = self.classify_sheet_with_model(
            self.bert_checkpoint, self.full_data
        )
        electra_sheet = self.classify_sheet_with_model(
            self.electra_checkpoint, self.full_data
        )

        # Save to CSV files instead of Excel
        bert_file = os.path.join(self.destination_path, f"SplitBert.csv")
        electra_file = os.path.join(self.destination_path, f"SplitElectra.csv")

        bert_sheet.to_csv(bert_file, index=False)
        electra_sheet.to_csv(electra_file, index=False)

        return bert_trainer, electra_trainer

    def inference_documents(
        self,
        input_folder: str,
        output_folder: str,
        threshold: float = 0.15,
        model_type: str = "bert",
    ):
        """
        Process all Word documents in a folder using the trained model.

        Args:
            input_folder (str): Path to folder containing .docx files
            output_folder (str): Path to save processed documents
            threshold (float): Confidence threshold for classification
            model_type (str): Model to use for inference ('bert' or 'electra')
        """

        if self.bert_checkpoint is None and self.electra_checkpoint is None:
            return ValueError(
                "You haven't trained the models yet , Please make sure to run the train_and_classify method first"
            )

        def generate_lower_colors(n):
            """Generate n distinct highlight colors from WD_COLOR_INDEX"""
            all_colors = [
                WD_COLOR_INDEX.BLUE,
                WD_COLOR_INDEX.BRIGHT_GREEN,
                WD_COLOR_INDEX.DARK_BLUE,
                WD_COLOR_INDEX.DARK_RED,
                WD_COLOR_INDEX.DARK_YELLOW,
                WD_COLOR_INDEX.GRAY_25,
                WD_COLOR_INDEX.GRAY_50,
                WD_COLOR_INDEX.GREEN,
                WD_COLOR_INDEX.PINK,
                WD_COLOR_INDEX.RED,
                WD_COLOR_INDEX.TEAL,
                WD_COLOR_INDEX.TURQUOISE,
                WD_COLOR_INDEX.VIOLET,
                WD_COLOR_INDEX.YELLOW,
            ]
            if n <= len(all_colors):
                return all_colors[:n]
            else:
                # If we need more colors than available, cycle through the colors
                return [all_colors[i % len(all_colors)] for i in range(n)]

        def generate_upper_colors(n):
            """Generate n distinct RGB colors using HSV color space"""
            colors = []
            for i in range(n):
                # Use golden ratio to space out hues evenly
                hue = i * (360 / n)
                # Convert HSV to RGB (using full saturation and value)
                h = hue / 360
                # Simple HSV to RGB conversion
                if h < 1 / 6:
                    r, g, b = 255, int(h * 6 * 255), 0
                elif h < 2 / 6:
                    r, g, b = int((2 / 6 - h) * 6 * 255), 255, 0
                elif h < 3 / 6:
                    r, g, b = 0, 255, int((h - 2 / 6) * 6 * 255)
                elif h < 4 / 6:
                    r, g, b = 0, int((4 / 6 - h) * 6 * 255), 255
                elif h < 5 / 6:
                    r, g, b = int((h - 4 / 6) * 6 * 255), 0, 255
                else:
                    r, g, b = 255, 0, int((1 - h) * 6 * 255)
                colors.append(RGBColor(r, g, b))
            return colors

        # Generate colors based on number of labels and top levels
        COLOR_LIST = generate_lower_colors(len(self.LABELS))
        TOP_LEVELS = list(self.upper_lower_mapping.keys())
        TOP_LEVEL_COLORS = generate_upper_colors(len(TOP_LEVELS))

        # Create color mappings
        low_color_dict = dict(zip(self.LABELS, COLOR_LIST))
        high_color_dict = dict(zip(TOP_LEVELS, TOP_LEVEL_COLORS))

        # Get top levels from upper_lower_mapping

        def apply_low_highlight(run, label):
            color_index = low_color_dict.get(label, WD_COLOR_INDEX.AUTO)
            run.font.highlight_color = color_index

        def apply_high_highlight(run, label):
            color = high_color_dict.get(label, None)
            # Ensure color is an instance of RGBColor; if not, use default black
            if not (
                hasattr(color, "rgb") or isinstance(color, type(RGBColor(0, 0, 0)))
            ):
                color = RGBColor(0, 0, 0)
            run.font.color.rgb = color

        # Load model checkpoint based on model_type
        if model_type.lower() == "electra":
            model_checkpoint = self.electra_checkpoint
        else:  # default to bert
            model_checkpoint = self.bert_checkpoint

        # Create output directory
        os.makedirs(output_folder, exist_ok=True)

        # Create classifier once outside the loop for efficiency
        classifier = pipeline("text-classification", model_checkpoint)

        print(f"Using {model_type.upper()} model for inference...")

        # Process each document
        for filename in tqdm.tqdm(os.listdir(input_folder)):
            if not filename.endswith(".docx"):
                continue

            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, f"processed_{filename}")

            doc = Document(input_path)

            high_label_counts = {label: 0 for label in TOP_LEVELS}
            low_label_counts = {label: 0 for label in self.LABELS}

            # Process paragraphs
            for paragraph in doc.paragraphs:
                if re.match(r"^\s*interviewee\s*:", paragraph.text, re.IGNORECASE):
                    interview_text = re.sub(
                        r"^\s*interviewee\s*:\s*",
                        "",
                        paragraph.text,
                        flags=re.IGNORECASE,
                    )

                    if len(interview_text.split()) > 10:
                        # Use the model pipeline to classify
                        prediction = classifier(interview_text)[0]

                        if prediction["score"] > threshold:
                            # Determine top level category
                            low_label = prediction["label"]
                            for (
                                top_level,
                                categories,
                            ) in self.upper_lower_mapping.items():
                                if low_label in categories:
                                    high_label = top_level
                                    break

                            high_label_counts[high_label] += 1
                            low_label_counts[low_label] += 1

                            # Clear and rewrite paragraph without Color Coding the label prefix
                            for run in paragraph.runs:
                                run.clear()

                            paragraph.add_run(
                                "Interviewee: "
                            )  # static text without highlighting
                            run = paragraph.add_run(
                                interview_text
                            )  # text to be highlighted
                            apply_high_highlight(run, high_label)
                            apply_low_highlight(run, low_label)

            # Add legends and summary
            legend_paragraph = doc.add_paragraph(
                "LEGEND: TOP LEVEL COLOR IDENTIFICATION"
            )
            for label, color_index in high_color_dict.items():
                run = legend_paragraph.add_run(f"\n{label}:")
                run.font.color.rgb = color_index

            legend_paragraph = doc.add_paragraph(
                "\nLEGEND: SUB LEVEL COLOR IDENTIFICATION"
            )
            for label, color_index in low_color_dict.items():
                run = legend_paragraph.add_run(f"\n{label}:")
                run.font.highlight_color = color_index

            # Add summary statistics
            summary_paragraph = doc.add_paragraph("\n\nSUMMARY:\n")

            summary_paragraph.add_run("\nHigh-level label counts:\n").bold = True
            for label, count in high_label_counts.items():
                summary_paragraph.add_run(f"{label}: {count}\n")

            summary_paragraph.add_run("\nLow-level label counts:\n").bold = True
            for label, count in low_label_counts.items():
                summary_paragraph.add_run(f"{label}: {count}\n")

            # Save processed document
            doc.save(output_path)
            print(f"Processed {filename}")


if __name__ == "__main__":
    trainer = TranscriptTrainer(
        input_file="CompletedMerged.csv",  # can be .xlsx or .csv
        destination_path="results",
        max_length=512,
        num_train_epochs=10,
    )

    bert_model, electra_model = trainer.train_and_classify()
