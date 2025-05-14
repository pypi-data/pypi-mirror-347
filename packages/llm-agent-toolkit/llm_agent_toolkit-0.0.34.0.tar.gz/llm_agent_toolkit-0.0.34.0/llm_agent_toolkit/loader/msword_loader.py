import os
import logging
import zipfile
from io import StringIO, BytesIO
import re
from contextlib import contextmanager

from docx import Document
from docx.document import Document as _Document
from docx.table import Table

from .._loader import BaseLoader
from .._core import ImageInterpreter
# from .._util import MessageBlock


logger = logging.getLogger(__name__)
"""
Dependencies:
----------
- python-docx==1.1.2
"""


class MsWordLoader(BaseLoader):
    def __init__(
        self,
        text_only: bool = True,
        tmp_directory: str | None = None,
        image_interpreter: ImageInterpreter | None = None,
    ):
        self.__text_only = text_only
        self.__image_interpreter = image_interpreter
        self.__tmp_directory = tmp_directory
        if not text_only:
            assert isinstance(tmp_directory, str)
            tmp_directory = tmp_directory.strip()
            if not tmp_directory:
                raise ValueError(
                    "Invalid temporary directory: Must be a non-empty string."
                )

            if not os.path.exists(tmp_directory):
                logger.warning(
                    "Temporary directory not exists. Will create one with name: %s",
                    tmp_directory,
                )
                os.makedirs(tmp_directory)

    @staticmethod
    def raise_if_invalid(input_path: str) -> None:
        if not all(
            [input_path is not None, isinstance(input_path, str), input_path != ""]
        ):
            raise ValueError("Invalid input path: Path must be a non-empty string.")

        if input_path[-5:] != ".docx":
            raise ValueError("Unsupported file format: Must be a DOCX file.")

        if not os.path.exists(input_path):
            raise FileNotFoundError(f"File not found: '{input_path}'.")

    def load(self, input_path: str) -> str:
        MsWordLoader.raise_if_invalid(input_path)

        markdown_content = []

        doc = Document(input_path)

        # Handle text content
        markdown_content.extend(self.extract_text_content(doc))

        # Handle tables
        markdown_content.extend(self.extract_tables_content(doc))

        # Handle images
        markdown_content.extend(self.extract_image_content(input_path))

        return "\n".join(markdown_content)

    async def load_async(self, input_path: str) -> str:
        MsWordLoader.raise_if_invalid(input_path)

        markdown_content = []

        doc = Document(input_path)

        # Handle text content
        markdown_content.extend(self.extract_text_content(doc))

        # Handle tables
        markdown_content.extend(self.extract_tables_content(doc))

        # Handle images
        markdown_content.extend(await self.extract_image_content_async(input_path))

        return "\n".join(markdown_content)

    @staticmethod
    def extract_text_content(doc: _Document) -> list[str]:
        markdown_content = []

        # Iterate through all elements in the document
        for para in doc.paragraphs:
            p = para._element
            pstyle_match = re.search(r'<w:pStyle w:val="([^"]+)"/>', p.xml)
            style_name = pstyle_match.group(1) if pstyle_match else "Normal"
            text = para.text.strip()
            if style_name == "Title":
                content = f"\n# {text}"
            elif style_name.startswith("Heading"):
                level = min(int(style_name[len("Heading") :]) + 1, 6)
                content = f"\n{'#' * level} {text}"
            elif style_name == "ListParagraph":
                content = f"* {text}"
            else:
                content = f"{text}"

            markdown_content.append(content)

        return markdown_content

    @staticmethod
    def extract_tables_content(doc: _Document) -> list[str]:
        """
        Iteratively extract tables from the document.

        Args:
            doc (_Document): The document to extract tables from.

        Returns:
            list[str]: A list of Markdown-formatted table content.

        Notes:
            - Ignore content formatting/styles.
            - Markdown syntax is used for the main table, nested tables are presented in HTML.
            - Exact location of the tables in the document is not guaranteed.
        """

        markdown_content = []

        # Extract tables
        for table_index, table in enumerate(doc.tables, start=1):
            markdown_content.append(f"## Table {table_index}\n")
            markdown_content.append(MsWordLoader.extract_table_content(table))
            # markdown_content.append("\n")

        if len(markdown_content) > 0:
            markdown_content.insert(0, "\n# Tables\n")

        return markdown_content

    @staticmethod
    def extract_table_content(table: Table):
        """
        Extract table content, support nested table.

        Args:
            table (Table): The table to extract content from.

        Returns:
            str: The extracted table content in Markdown format.

        Notes:
            - Ignore content formatting/styles.
            - Markdown syntax is used for the main table, nested tables are presented in HTML.
        """
        markdown_content = []

        header_row = table.rows[0]
        headers = [f"| {cell.text.strip()} " for cell in header_row.cells]
        markdown_content.append("".join(headers) + "|")
        markdown_content.append(
            f"|{'---|' * len(headers)}"
        )  # Markdown header separator

        for row in table.rows[1:]:
            row_content = []
            for cell in row.cells:
                if cell.tables:  # Check if the cell contains a nested table
                    nested_table_md = []
                    for nested_table in cell.tables:
                        nested_table_md.append(
                            MsWordLoader.extract_subtable_content(nested_table)
                        )
                    row_content.append(
                        f"{cell.text.strip()} {''.join(nested_table_md)}"
                    )
                else:
                    row_content.append(f"{cell.text.strip()}")
            row_string = f"| {' | '.join(row_content)} |"
            markdown_content.append(row_string)

        return "\n".join(markdown_content)

    @staticmethod
    def extract_subtable_content(table: Table) -> str:
        """
        Recursively extract table content.

        Args:
            table (Table): The table to extract content from.

        Returns:
            str: The extracted content.

        Notes:
            - This function is used to extract content from nested tables.
            - It is called recursively for each nested table.
            - Ignore content formatting/styles.
            - HTML table structure is used to represent the extracted content.
        """
        content = StringIO()
        content.write("<table>")
        for row in table.rows:
            content.write("<tr>")
            for cell in row.cells:
                content.write(f"<td>{cell.text.strip()}")
                if cell.tables:
                    for nested_table in cell.tables:
                        content.write(
                            MsWordLoader.extract_subtable_content(nested_table)
                        )
                content.write("</td>")
            content.write("</tr>")
        content.write("</table>")
        return content.getvalue()

    # @staticmethod
    # def get_cell_content_with_formatting(cell: _Cell):
    #     """
    #     Extracts the content of a cell with formatting.

    #     Notes:
    #     * Tables in
    #     """
    #     content = StringIO()
    #     for para in cell.paragraphs:
    #         for run in para.runs:
    #             if run.bold:
    #                 content.write(f"**{run.text}**")
    #             elif run.italic:
    #                 content.write(f"*{run.text}*")
    #             else:
    #                 content.write(run.text)
    #         content.write("\n")
    #     return content.getvalue().strip()

    @staticmethod
    def extract_alt_text_dict(docx: zipfile.ZipFile) -> dict[str, str]:
        """
        Extracts alt text for images in a DOCX file.

        Parameters:
        ----------
        - docx: zipfile.ZipFile: The ZipFile object representing the DOCX file

        Returns:
        ----------
        * dict[str, str]: Dictionary mapping image file names to their alt text

        """
        from xml.etree import ElementTree as ET

        image_alt_texts = {}
        # Parse the XML document to get image alt text
        if "word/document.xml" in docx.namelist():
            document_xml = docx.read("word/document.xml")
            root = ET.fromstring(document_xml)

            # Find all elements with cNvPr to identify inserted images

            for elem in root.iter():
                if elem.tag.endswith("cNvPr"):
                    r_id = elem.attrib.get("id")
                    alt_text = elem.attrib.get("descr", "Alt text not available")
                    ele_name = elem.attrib.get("name", str(r_id))  # Picture {index}
                    if ele_name and alt_text:
                        image_alt_texts[ele_name] = (
                            alt_text  # Add alt text if available
                        )

        return image_alt_texts

    @staticmethod
    def get_alt_by_name(d: dict[str, str], key1: str, key2: str) -> str:
        # This is needed because `extract_alt_text_dict` keys follow the pattern `Picture {index}` or `{index}`
        # key1: `Picture {index}`
        # key2: `{index}`
        return d.get(key1, d.get(key2, "Alt text not available"))

    def extract_image_content(self, input_path: str) -> list[str]:
        markdown_content = []

        with zipfile.ZipFile(input_path, "r") as docx:
            image_alt_texts: dict[str, str] = self.extract_alt_text_dict(docx)

            # Iterate through the files in the archive
            file_startswith_word_media_lst = list(
                filter(lambda f: f.startswith("word/media/"), docx.namelist())
            )
            for counter, file in enumerate(file_startswith_word_media_lst, start=1):
                # Extract the corresponding alt text
                # Assumption: Images are captured in the same order as `extract_alt_text_dict`
                alt_text = self.get_alt_by_name(
                    image_alt_texts, key1=f"Picture {counter}", key2=str(counter)
                )

                if self.__text_only or self.__image_interpreter is None:
                    image_description = "Image description not available"
                else:
                    image_data = docx.read(file)
                    image_name = os.path.basename(file)  # image{index}.png
                    with self.temporary_file(image_data, image_name) as image_path:
                        responses, usage = self.__image_interpreter.interpret(
                            query="Describe this image",
                            context=None,
                            filepath=image_path,
                        )
                        image_description = responses[0]["content"]

                markdown_content.append(
                    f"## {os.path.basename(file)}\nDescription: {image_description}\n\nAlt Text: {alt_text}\n"
                )
        if len(markdown_content) > 0:
            markdown_content.insert(0, "\n# Images\n")

        return markdown_content

    async def extract_image_content_async(self, input_path: str) -> list[str]:
        markdown_content = []

        with zipfile.ZipFile(input_path, "r") as docx:
            image_alt_texts: dict[str, str] = self.extract_alt_text_dict(docx)

            # Iterate through the files in the archive
            file_startswith_word_media_lst = list(
                filter(lambda f: f.startswith("word/media/"), docx.namelist())
            )
            for counter, file in enumerate(file_startswith_word_media_lst, start=1):
                # Extract the corresponding alt text
                # Assumption: Images are captured in the same order as `extract_alt_text_dict`
                alt_text = self.get_alt_by_name(
                    image_alt_texts, key1=f"Picture {counter}", key2=str(counter)
                )

                if self.__text_only or self.__image_interpreter is None:
                    image_description = "Image description not available"
                else:
                    image_data = docx.read(file)
                    image_name = os.path.basename(file)  # image{index}.png
                    with self.temporary_file(image_data, image_name) as image_path:
                        (
                            responses,
                            usage,
                        ) = await self.__image_interpreter.interpret_async(
                            query="Describe this image",
                            context=None,
                            filepath=image_path,
                        )
                        image_description = responses[0]["content"]

                markdown_content.append(
                    f"## {os.path.basename(file)}\nDescription: {image_description}\n\nAlt Text: {alt_text}\n"
                )
        if len(markdown_content) > 0:
            markdown_content.insert(0, "\n# Images\n")

        return markdown_content

    @contextmanager
    def temporary_file(self, image_bytes: bytes, filename: str):
        tmp_path = f"{self.__tmp_directory}/{filename}"
        try:
            image_stream = BytesIO(image_bytes)
            image_stream.seek(0)
            with open(tmp_path, "wb") as f:
                f.write(image_bytes)
            yield tmp_path
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    @contextmanager
    def increment_later(self, counter: int):
        yield counter
        return counter + 1
