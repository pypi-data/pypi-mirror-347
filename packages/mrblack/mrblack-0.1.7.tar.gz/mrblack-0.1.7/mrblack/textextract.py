#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# File: textextract.py
# Author: Wadih Khairallah
# Created: 2024-12-01 12:12:08
# Modified: 2025-05-15 03:03:15

import os
import re
import json
import math
import socket
import platform
import subprocess
import hashlib
import string
import tempfile
import random
import unicodedata


from uuid import uuid4
from datetime import datetime
from io import StringIO
from typing import (
    Optional,
    Dict,
    Any,
    List,
)

# Web Specific
import requests
from requests_html import HTMLSession
from urllib.parse import urlparse
from bs4 import BeautifulSoup

# Document Specific
from readability import Document as RDocument
import magic
import pytesseract
import pandas as pd
import speech_recognition as sr
import pymupdf
from docx import Document
from mss import mss
from pydub import AudioSegment

# Image Specific
from PIL import Image

from collections import Counter
from rich.console import Console

console = Console()
print = console.print
log = console.log


USER_AGENTS = [
    # Desktop
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 13_3) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.0 Safari/605.1.15",
    "Mozilla/5.0 (X11; Linux x86_64) Gecko/20100101 Firefox/115.0",
    # Mobile
    "Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.0 Mobile/15E148 Safari/604.1",
    "Mozilla/5.0 (Linux; Android 13; Pixel 7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.6367.91 Mobile Safari/537.36"
]


def clean_path(
    path: str
) -> Optional[str]:
    """
    Normalize and validate a filesystem path.

    Args:
        path (str): Input file or directory path.

    Returns:
        Optional[str]: Absolute path if valid; None otherwise.
    """
    p = os.path.expanduser(path)
    p = os.path.abspath(p)
    if os.path.isfile(p) or os.path.isdir(p):
        return p
    return None


def normalize(
    text: str
) -> str:
    """
    Replace multiple consecutive newlines, carriage returns, and spaces
    with a single space. Ensures compact, single-line output.

    Args:
        text (str): Raw input text.

    Returns:
        str: Normalized single-line text.
    """
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'(?m)(^ \n)+', '\n', text)
    text = re.sub(r'\t+', '\t', text)
    text = re.sub(r'\r+', '\n', text)
    text = re.sub(r"^ ", "", text, flags=re.MULTILINE)
    return text 


def is_url(s: str) -> bool:
    parsed = urlparse(s)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def text_from_screenshot() -> str:
    """
    Capture a full-screen screenshot, perform OCR, and clean up temp file.

    Returns:
        str: Normalized OCR-extracted text from the screenshot.
    """
    from uuid import uuid4
    tmp_filename = f"screenshot_{uuid4().hex}.png"
    tmp_path = os.path.join(tempfile.gettempdir(), tmp_filename)

    try:
        with mss() as sct:
            monitor = {"top": 0, "left": 0, "width": 0, "height": 0}
            for mon in sct.monitors:
                monitor["left"] = min(mon["left"], monitor["left"])
                monitor["top"] = min(mon["top"], monitor["top"])
                monitor["width"] = max(mon["width"] + mon["left"] - monitor["left"], monitor["width"])
                monitor["height"] = max(mon["height"] + mon["top"] - monitor["top"], monitor["height"])
            screenshot = sct.grab(monitor)

        img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")
        img_gray = img.convert("L")
        img_gray.save(tmp_path)

        content = text_from_image(tmp_path)
        return normalize(content)
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception as e:
                print(f"Failed to delete temp screenshot: {e}")


def extract_exif(
    file_path: str
) -> Optional[Dict[str, Any]]:
    """
    Extract EXIF metadata from a file using exiftool.

    Args:
        file_path (str): Path to the target file.

    Returns:
        Optional[Dict[str, Any]]: Parsed EXIF data, or None on failure.
    """
    exif_data: Optional[Dict[str, Any]] = None
    try:
        result = subprocess.run(
            ['exiftool', '-j', file_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        if result.returncode == 0:
            exif_data = json.loads(result.stdout.decode())[0]
    except Exception as e:
        print(f"Exiftool failed: {e}")
    return exif_data


def text_from_html(html: str) -> str:
    """
    Extract readable text from raw HTML content.

    Args:
        html (str): HTML source as a string.

    Returns:
        str: Cleaned and normalized visible text.
    """
    soup = BeautifulSoup(html, "html.parser")

    # Remove non-visible or structural elements
    for tag in soup([
        "script", "style",
        "noscript", "iframe",
        "meta", "link",
        "header", "footer",
        "form", "nav",
        "aside"
    ]):
        tag.decompose()

    text = soup.get_text(separator=" ")

    return normalize(text)


def text_from_url(
    url: str,
    render_js: bool = True
) -> Optional[str]:
    """
    Fetch and extract all visible text from a web page, including JS-rendered content.

    Args:
        url (str): Target webpage URL.
        render_js (bool): Whether to render JavaScript content.

    Returns:
        Optional[str]: Cleaned full-page text, or None on failure.
    """
    headers = {
        "User-Agent": random.choice(USER_AGENTS),
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": url,
        "DNT": "1",
        "Upgrade-Insecure-Requests": "1"
    }

    session = HTMLSession()
    try:
        r = session.get(url, headers=headers, timeout=20)
        if render_js:
            r.html.render(timeout=20, sleep=1)

        html = r.html.html
        content = text_from_html(html)

        return content

    except Exception as e:
        print(f"[Error] {url} - {e}")
        return None
    finally:
        session.close()


def extract_text(
    file_path: str
) -> Optional[str]:
    """
    Extract text content from a local file or URL.

    Supports web pages, text, JSON, XML, CSV, Excel, PDF, DOCX, images, audio.

    Args:
        file_path (str): Path to the input file or URL.

    Returns:
        Optional[str]: Extracted text, or None if unsupported or error.
    """
    if is_url(file_path):
        return text_from_url(file_path)

    TEXT_MIME_TYPES = {
        "application/json", "application/xml", "application/x-yaml",
        "application/x-toml", "application/x-csv", "application/x-markdown",
    }

    path = clean_path(file_path)
    if not path:
        print(f"No such file: {file_path}")
        return None

    mime_type = magic.from_file(path, mime=True)
    try:
        if mime_type.startswith("text/html"):
            content = text_from_html(path)

        if mime_type.startswith("text/") or mime_type in TEXT_MIME_TYPES:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

        elif mime_type in [
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ]:
            content = text_from_excel(path)

        elif mime_type == "application/pdf":
            content = text_from_pdf(path)

        elif mime_type == \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            content = text_from_docx(path)

        elif mime_type == "application/msword":
            content = text_from_doc(path)

        elif mime_type.startswith("image/"):
            content = text_from_image(path)

        elif mime_type.startswith("audio/"):
            content = text_from_audio(path)

        else:
            content = text_from_any(path)

        if content:
            return content
        else:
            print(f"No content found for file: {path}")
            return None
    except Exception as e:
        print(f"Error reading {path}: {e}")
        return None

def text_from_audio(
    audio_file: str
) -> Optional[str]:
    """
    Transcribe audio to text using Google Speech Recognition.

    Args:
        audio_file (str): Path to the input audio file.

    Returns:
        Optional[str]: Transcription, or None on failure.
    """
    def convert_to_wav(file_path: str) -> str:
        _, ext = os.path.splitext(file_path)
        ext = ext.lstrip('.')
        audio = AudioSegment.from_file(file_path, format=ext)
        tmp_filename = f"audio_{uuid4().hex}.wav"
        wav_path = os.path.join(tempfile.gettempdir(), tmp_filename)
        audio.export(wav_path, format='wav')
        return wav_path

    recognizer = sr.Recognizer()
    temp_wav_path = None
    cleanup_needed = False

    try:
        _, ext = os.path.splitext(audio_file)
        if ext.lower() not in ['.wav', '.wave']:
            temp_wav_path = convert_to_wav(audio_file)
            cleanup_needed = True
        else:
            temp_wav_path = clean_path(audio_file)

        if not temp_wav_path:
            print("Invalid audio path.")
            return None

        with sr.AudioFile(temp_wav_path) as source:
            audio = recognizer.record(source)
        return recognizer.recognize_google(audio)

    except sr.UnknownValueError:
        print("Could not understand audio.")
    except sr.RequestError as e:
        print(f"Speech recognition error: {e}")
    except Exception as e:
        print(f"Failed to process audio: {e}")
    finally:
        if cleanup_needed and temp_wav_path and os.path.exists(temp_wav_path):
            try:
                os.remove(temp_wav_path)
            except Exception as e:
                print(f"Failed to delete temp WAV file {temp_wav_path}: {e}")

    return None


def downloadImage(
    url: str
) -> Optional[str]:
    """
    Download an image from a URL to /tmp/ and return its path.

    Args:
        url (str): Remote image URL.

    Returns:
        Optional[str]: Local file path, or None on failure.
    """
    if is_image(url):
        filename = os.path.basename(urlparse(url).path)
        save_path = os.path.join("/tmp/", filename)
        resp = requests.get(url, stream=True)
        resp.raise_for_status()
        with open(save_path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)
        return clean_path(save_path)
    print(f"Unable to pull image from {url}")
    return None


def is_image(
    file_path_or_url: str
) -> bool:
    """
    Determine if the given path/URL points to an image.

    Args:
        file_path_or_url (str): Local path or URL.

    Returns:
        bool: True if MIME type starts with 'image/'.
    """
    try:
        mime = magic.from_file(file_path_or_url, mime=True)
        return mime.startswith("image/")
    except Exception:
        return False


def text_from_pdf(
    pdf_path: str
) -> Optional[str]:
    """
    Extract text and OCR results from a PDF using PyMuPDF.

    Args:
        pdf_path (str): Path to PDF file.

    Returns:
        Optional[str]: Combined normalized text and image OCR results.
    """
    plain_text = ""
    temp_image_paths: List[str] = []

    try:
        doc = pymupdf.open(pdf_path)
        for k, v in doc.metadata.items():
            plain_text += f"{k}: {v}\n"

        for i in range(len(doc)):
            page = doc.load_page(i)
            plain_text += f"\n--- Page {i + 1} ---\n"
            text = page.get_text()
            plain_text += text or "[No text]\n"

            for img_index, img in enumerate(page.get_images(full=True), start=1):
                xref = img[0]
                base = doc.extract_image(xref)
                img_bytes = base["image"]

                img_filename = f"pdf_page{i+1}_img{img_index}_{uuid4().hex}.png"
                img_path = os.path.join(tempfile.gettempdir(), img_filename)
                temp_image_paths.append(img_path)

                with open(img_path, "wb") as f:
                    f.write(img_bytes)

                ocr = text_from_image(img_path) or ""
                plain_text += f"\n[Image {img_index} OCR]\n{ocr}\n"

        return normalize(plain_text)
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return None
    finally:
        for path in temp_image_paths:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except Exception as e:
                    print(f"Failed to delete temp image {path}: {e}")
        doc.close()


def text_from_doc(
    filepath: str,
    min_length: int = 4
) -> str:
    """
    Extract readable strings and metadata from binary Word (.doc) files.

    Args:
        filepath (str): Path to .doc file.
        min_length (int): Minimum string length to extract.

    Returns:
        str: Metadata and text content.
    """
    def extract_printable_strings(
        data: bytes
    ) -> List[str]:
        pattern = re.compile(
            b'[' + re.escape(bytes(string.printable, 'ascii')) +
            b']{%d,}' % min_length
        )
        found = pattern.findall(data)
        return list(dict.fromkeys(m.decode(errors='ignore').strip()
                                   for m in found))

    def clean_strings(
        strs: List[str]
    ) -> List[str]:
        cleaned: List[str] = []
        skip = ["HYPERLINK", "OLE2", "Normal.dotm"]
        for line in strs:
            if any(line.startswith(pref) for pref in skip):
                continue
            cleaned.append(re.sub(r'\s+', ' ', line).strip())
        return cleaned

    with open(filepath, 'rb') as f:
        data = f.read()
    strings = extract_printable_strings(data)
    strings = clean_strings(strings)
    content = "\n".join(strings)
    return normalize(content)


def text_from_docx(
    file_path: str
) -> Optional[str]:
    """
    Extract text, tables, and OCR from embedded images in a DOCX file.

    Args:
        file_path (str): Path to the .docx file.

    Returns:
        Optional[str]: Normalized full text content.
    """
    path = clean_path(file_path)
    if not path:
        return None

    temp_image_paths: List[str] = []
    plain_text = ""

    try:
        doc = Document(path)

        for p in doc.paragraphs:
            if p.text.strip():
                plain_text += p.text.strip() + "\n"

        for tbl in doc.tables:
            plain_text += "\n[Table]\n"
            for row in tbl.rows:
                row_text = "\t".join(c.text.strip() for c in row.cells)
                plain_text += row_text + "\n"

        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                blob = rel.target_part.blob

                img_filename = f"docx_img_{rel_id}_{uuid4().hex}.png"
                img_path = os.path.join(tempfile.gettempdir(), img_filename)
                temp_image_paths.append(img_path)

                with open(img_path, "wb") as img_file:
                    img_file.write(blob)

                ocr = text_from_image(img_path) or ""
                plain_text += f"\n[Image OCR]\n{ocr}\n"

        return normalize(plain_text)

    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return None
    finally:
        for path in temp_image_paths:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except Exception as e:
                    print(f"Failed to delete temp DOCX image {path}: {e}")


def text_from_excel(
    file_path: str
) -> str:
    """
    Convert an Excel workbook to CSV text.

    Args:
        file_path (str): Path to the Excel file.

    Returns:
        str: CSV-formatted string.
    """
    path = clean_path(file_path)
    if not path:
        return ""
    try:
        df = pd.read_excel(path)
        out = StringIO()
        df.to_csv(out, index=False)
        return out.getvalue()
    except Exception as e:
        print(f"Failed Excel -> CSV: {e}")
        return ""


def text_from_image(
    file_path: str
) -> Optional[str]:
    """
    Perform OCR on an image file.

    Args:
        file_path (str): Path to the image.

    Returns:
        Optional[str]: Extracted text, or None on error.
    """
    path = clean_path(file_path)
    if not path:
        return None
    try:
        with Image.open(path) as img:
            txt = pytesseract.image_to_string(img).strip()
            return normalize(txt) or ""
    except Exception as e:
        print(f"Failed image OCR: {e}")
        return None


def text_from_any(
    file_path: str
) -> Optional[str]:
    """
    Handle unknown file types by reporting stats and metadata.

    Args:
        file_path (str): Path to the file.

    Returns:
        Optional[str]: Plain-text report, or None on error.
    """
    path = clean_path(file_path)
    if not path:
        return None
    try:
        stats = os.stat(path)
        info = {
            "path": path,
            "size": stats.st_size,
            "created": datetime.fromtimestamp(stats.st_ctime).isoformat(),
            "modified": datetime.fromtimestamp(stats.st_mtime).isoformat(),
        }
        content = "\n".join(f"{k}: {v}" for k, v in info.items())
        return normalize(content)
    except Exception as e:
        print(f"Error on other file: {e}")
        return None


def extract_metadata(
    file_path: str
) -> Dict[str, Any]:
    """
    Extract comprehensive metadata from any file type.

    Args:
        file_path (str): Path to target file.

    Returns:
        Dict[str, Any]: Nested metadata structure.
    """
    path = clean_path(file_path)
    if not path:
        return {"error": "File not found"}
    meta: Dict[str, Any] = {}
    try:
        stats = os.stat(path)
        meta["size_bytes"] = stats.st_size
        meta["mime"] = magic.from_file(path, mime=True)
        meta["hashes"] = {
            "md5": hashlib.md5(open(path,'rb').read()).hexdigest()}
    except Exception as e:
        meta["error"] = str(e)
    return meta


def main() -> None:
    """
    CLI entry point for text or metadata extraction.
    Parses arguments and prints results.
    """
    import argparse
    parser = argparse.ArgumentParser(
        description="Extract text or metadata from any file or url"
    )
    parser.add_argument(
        "file",
        type=str,
        help="Path to the input file"
    )
    parser.add_argument(
        "--metadata",
        action="store_true",
        help="Extract metadata instead of text"
    )
    args = parser.parse_args()
    if args.metadata:
        data = extract_metadata(args.file)
        print(json.dumps(data, indent=2))
    else:
        txt = extract_text(args.file)
        print(txt or "No text extracted.")


if __name__ == "__main__":
    main()

